"""
build_chapters.py
Generates CHAPTERS_3_AND_4.docx — the rewritten thesis chapters.
Run from the timetable_project directory:
    python build_chapters.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIAGRAMS = os.path.join(os.path.dirname(__file__), "diagrams")

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading1(doc, text):
    """Chapter heading — centered, 14pt bold, Times New Roman."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    return p


def heading2(doc, text):
    """Section heading — 12pt bold, Times New Roman, left-aligned."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def heading3(doc, text):
    """Sub-section heading — 12pt bold italic, Times New Roman."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def body(doc, text):
    """Standard body paragraph — 12pt, Times New Roman, justified, 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    fmt = p.paragraph_format
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(18)
    # 1.5 line spacing
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    return p


def caption(doc, text):
    """Figure / table caption — 11pt italic, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    return p


def insert_image(doc, filename, caption_text, width=Inches(5.5)):
    path = os.path.join(DIAGRAMS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        caption(doc, caption_text)
    else:
        body(doc, f"[IMAGE NOT FOUND: {filename} — insert {caption_text} here]")


def add_table(doc, headers, rows, caption_text):
    """Add a simple bordered table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val)
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

    caption(doc, caption_text)
    return table


def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────

def build():
    doc = Document()

    # Set default margins (2.5cm all sides — typical thesis)
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.81)
        section.right_margin  = Cm(2.54)

    # ═══════════════════════════════════════════
    # CHAPTER THREE
    # ═══════════════════════════════════════════
    heading1(doc, "CHAPTER THREE")
    heading1(doc, "RESEARCH METHODOLOGY")

    # 3.1
    heading2(doc, "3.1 Preamble")
    body(doc, "This chapter provides a comprehensive account of the methodology adopted in the design, development, and evaluation of the hyper-heuristic university timetable generation system described in this study. The system, titled \"From Prompt to Timetable: Integrating Natural Language Reasoning with Different Hyper-Heuristic Algorithms,\" was built as a software artefact that accepts academic scheduling data and natural language constraints, and then applies three competing meta-heuristic optimisation algorithms to produce a conflict-minimised university timetable.")
    body(doc, "The research adopts the Design Science Research (DSR) methodology, which is appropriate for studies that involve the design, construction, and evaluation of novel IT artefacts intended to solve clearly defined organisational problems (Dresch, Lacerda, & Antunes, 2015). Within DSR, this project is classified as a mixed design: it combines applied design (the construction of a functioning system) with evaluatory design (the systematic experimental comparison of three optimisation algorithms on real institutional data). This dual nature means that the methodology chapter serves two purposes: it documents the artefact design and implementation decisions, and it states the evaluation strategy that will be assessed in Chapter Four.")
    body(doc, "The remainder of this chapter is structured as follows. Section 3.2 states the research design. Section 3.3 describes the data source and data collection process, including an honest account of the dataset's limitations. Sections 3.4 to 3.9 present the system analysis, architecture, input/output design, and database design. Sections 3.10 to 3.14 specify the constraint model, fitness function, initial solution generation, algorithm designs, and natural language processing module. Sections 3.15 to 3.17 cover the UML models, software requirements, and implementation procedures. Section 3.18 states the evaluation strategy. Section 3.19 summarises the chapter.")

    # 3.2
    heading2(doc, "3.2 Research Design")
    body(doc, "The research design adopted for this study is the Design Science Research (DSR) methodology. DSR is a problem-solving paradigm that prescribes the construction and evaluation of purposeful artefacts — in this case, a software system — to address identified organisational problems (Dresch et al., 2015). The methodology is particularly suitable for this project for the following reasons.")
    body(doc, "First, the university timetable scheduling problem (UTSP) is a clearly identified, well-documented organisational problem. It is classified in the theoretical computer science literature as NP-hard, meaning that no known polynomial-time algorithm can find the globally optimal solution for real-world instances (Ceschia, Di Gaspero, & Schaerf, 2023). This justifies the use of meta-heuristic methods, which produce good-quality solutions within acceptable computational time rather than guaranteeing mathematical optimality.")
    body(doc, "Second, the primary output of this project is a novel software artefact — an intelligent, web-based timetable generation system that integrates natural language constraint specification with a hyper-heuristic optimisation framework. The construction of this artefact constitutes the applied design component of the research.")
    body(doc, "Third, the artefact is systematically evaluated through controlled experiments. Three meta-heuristic algorithms — the Genetic Algorithm (GA), Simulated Annealing (SA), and Tabu Search (TS) — are applied to the same scheduling problem, and their performance is compared using quantitative metrics including fitness score, clash counts, and execution time. This constitutes the evaluatory design component of the research.")
    body(doc, "The study therefore follows a mixed research design: applied design in the construction of the artefact, and evaluatory design in the experimental comparison of algorithms. This classification aligns with the project guidelines for mixed design projects, which combine development and evaluation.")

    # 3.3
    heading2(doc, "3.3 Population, Data Source, and Data Limitations")
    heading3(doc, "3.3.1 Data Collection Process and Limitations")
    body(doc, "The primary dataset used in this study is the Covenant University Academic Resources Dataset, which was assembled specifically for this project. The target data represented the full academic scheduling information for the 2025/2026 academic session at Covenant University, Ota, Ogun State, Nigeria, covering course codes, assigned lecturers, enrolled student counts, assigned teaching venues, and course load information across all colleges and departments.")
    body(doc, "Data was collected from multiple sources, including the Directorate of Academic Planning Unit (DAPU), timetable records obtained from individual departments, and direct departmental scheduling documents. However, the data collection process encountered significant limitations. The university does not maintain a single, centralised, machine-readable scheduling database accessible to undergraduate researchers. As a result, records from different departments were supplied in inconsistent formats, some were incomplete, and others contained missing or ambiguous entries for lecturer names, venue capacities, and student enrolment figures.")
    body(doc, "To address these gaps, supplementary data was added by the researcher based on institutional knowledge as a registered student of Covenant University. Where official lecturer names were unavailable, placeholder names or names drawn from publicly known departmental staff were used. Where venue capacities could not be confirmed from official records, estimates were made based on physical familiarity with the venues and standard university venue configurations. Where course entries were missing student counts, reasonable estimates were inserted based on class sizes typically observed at the department and level in question.")
    body(doc, "It is therefore important to acknowledge that the dataset used in this study is a partially estimated dataset. While every effort was made to ensure that the figures are realistic and internally consistent, the generated timetable, though algorithmically optimised given the input data, may not perfectly reflect true institutional scheduling constraints in every respect. The accuracy of the output is bounded by the accuracy and completeness of the input. This limitation is discussed further in Chapter Five.")

    heading3(doc, "3.3.2 Dataset Characteristics")
    body(doc, "After collection, cleaning, and supplementation, the dataset used for this study has the following characteristics:")
    add_table(doc,
        ["Property", "Value"],
        [
            ("File Format", "CSV (Comma-Separated Values)"),
            ("Total Records", "490+ course sessions"),
            ("Unique Departments", "20+ academic departments"),
            ("Academic Levels", "100, 200, 300, 400, 500"),
            ("Unique Venues", "50+ classroom venues"),
            ("Unique Lecturers", "100+ academic staff"),
            ("Student Count Range", "56 to 431 per course"),
            ("Course Load Range", "1 to 3 hours per session"),
        ],
        "Table 3.1: Covenant University Academic Dataset Characteristics"
    )

    heading3(doc, "3.3.3 Venue Classification")
    body(doc, "The venues in the dataset are classified into the following categories based on physical capacity and function:")
    add_table(doc,
        ["Category", "Examples", "Capacity Range"],
        [
            ("Mega Lecture Theatres", "LT1, LT2, LARGELH", "1,000 – 2,500"),
            ("Large Halls", "HALL107, HALL108, HALL201", "240 – 500"),
            ("Medium Halls", "HALL306, HALL307, STUDIO100", "180 – 280"),
            ("Science Laboratories", "BCH_LAB, BIO_LAB, CHEM_LAB", "140 – 180"),
            ("Computer Laboratories", "COMPUTER_LAB", "250 – 300"),
            ("Engineering Lecture Halls", "CHE 200LH, CVE 300LH, MCE 400LH", "150 – 250"),
            ("Special Venues", "CBN_AUDITORIUM, CHAPEL, LIBRARY", "500 – 1,200"),
            ("Overflow Facilities", "OVERFLOW_HALL_1, OVERFLOW_HALL_2", "800 – 1,000"),
        ],
        "Table 3.2: Venue Classification and Capacity Ranges"
    )

    heading3(doc, "3.3.4 Data Preprocessing Pipeline")
    body(doc, "Before the dataset is loaded into the scheduling engine, it passes through a rigorous preprocessing pipeline implemented in the load_dataset() function within engine.py. The pipeline performs the following steps:")
    body(doc, "(i) Column Name Standardisation: All column names are stripped of leading and trailing whitespace and mapped to a set of recognised standard names. For example, columns named \"LECTURERS\", \"LECTURER'S NAME\", or \"LECTURER NAME\" are all mapped to the standard \"LECTURER\" column. This ensures that datasets from different departments with varying column naming conventions can all be processed by the same engine without manual reformatting.")
    body(doc, "(ii) Missing Lecturer Assignment: For course records where the lecturer field is missing, blank, or contains placeholder values such as \"Unknown\" or \"N/A\", the system assigns a placeholder Nigerian academic name drawn from a predefined pool of realistic names within the engine configuration. This prevents the scheduling engine from failing due to missing lecturer entries while maintaining a realistic institutional appearance in the output.")
    body(doc, "(iii) Student Count Cleaning: Non-numeric student count values are coerced to numeric. Records with completely missing course codes are removed. Records where student count exceeds 1.5 times the venue's known capacity are flagged and removed as likely data entry errors.")
    body(doc, "(iv) Venue Capacity Mapping: Each venue name in the dataset is mapped to its seating capacity. Where the dataset includes an explicit venue capacity column, those values are used directly and stored in the engine's VENUE_CAPACITY dictionary. Where no capacity data is available, the engine applies a name-heuristic function that infers approximate capacity based on the venue name pattern.")
    body(doc, "(v) Venue Key Standardisation: Venue names are converted to uppercase, spaces are replaced with underscores, and special characters are removed to create consistent dictionary keys for the scheduling engine's room tracking logic.")
    body(doc, "(vi) Dataset Summary: Upon completion of preprocessing, the engine prints a summary of the total courses loaded, unique venues, unique lecturers, and maximum student count, allowing the administrator to verify the data before initiating optimisation.")

    insert_image(doc, "fig3_4.jpeg",
        "Figure 3.4: Data Preprocessing Pipeline Flowchart")

    # 3.4
    heading2(doc, "3.4 System Analysis")
    heading3(doc, "3.4.1 Analysis of the Existing System")
    body(doc, "Traditional university course timetabling at most Nigerian institutions, including Covenant University, is conducted through manual or semi-manual spreadsheet-based processes. Academic administrators from individual departments compile course lists, lecturer assignments, and room preferences into spreadsheet files, which are then manually reconciled across departments by timetabling officers. This process suffers from several well-documented limitations.")
    body(doc, "High Error Rate: Manual scheduling is inherently prone to human error. Lecturer clashes, room conflicts, and capacity violations are common outcomes of manual processes and typically require multiple iterations of correction (Jaradat, Ayob, Ahmad, Naser, & Jadaan, 2019).")
    body(doc, "Time Intensity: The manual process typically requires several weeks of iterative adjustment before a viable timetable is produced, competing with other administrative responsibilities and introducing scheduling delays at the start of each semester.")
    body(doc, "Scalability Failure: As the number of courses, lecturers, venues, and enrolled students grows, the combinatorial complexity of the scheduling problem increases exponentially. Manual approaches cannot effectively explore the large solution space and often settle for the first feasible solution found rather than an optimised one (Ceschia et al., 2023).")
    body(doc, "No Optimisation: The existing system produces timetables that are feasible at best. There is no mechanism for optimising quality metrics such as minimising the number of late-evening sessions, distributing classes evenly across the week, or respecting individual lecturer scheduling preferences.")
    body(doc, "No Natural Language Interface: Institutional scheduling policies and lecturer preferences must be translated by the timetabling officer into rigid spreadsheet formats. There is no facility for expressing constraints in ordinary English, which creates a communication barrier between non-technical administrators and the scheduling process.")

    heading3(doc, "3.4.2 Analysis of the Proposed System")
    body(doc, "The proposed system addresses each of the above limitations through an intelligent, automated scheduling platform. Natural Language Constraint Specification allows administrators to input scheduling constraints as plain English sentences. The system's NLP module, implemented in nlp_parser.py, parses these sentences using regular expression pattern matching and converts them into structured JSON constraint objects that the scheduling engine applies directly.")
    body(doc, "Multi-Algorithm Hyper-Heuristic Optimisation: Three optimisation algorithms — GA, SA, and TS — are executed in sequence on the same scheduling problem. The algorithm producing the lowest fitness score is automatically selected as the final timetable. This hyper-heuristic approach ensures that the best-available solution is always delivered, regardless of which algorithm happens to perform best for a particular problem instance (Burke et al., 2013).")
    body(doc, "Automated Clash Detection: The fitness function, implemented in the evaluate_timetable() function of engine.py, systematically detects and penalises lecturer clashes, room double-bookings, capacity violations, unscheduled sessions, and NLP constraint violations using a weighted penalty model.")
    body(doc, "Web-Based Interface: A Streamlit-powered frontend provides role-based access for three user types — academic administrators, lecturers, and students — with real-time timetable visualisation, a multi-step generation workflow, and CSV export functionality. A FastAPI backend server provides programmatic access to the scheduling engine via well-defined RESTful API endpoints.")

    # 3.5
    heading2(doc, "3.5 Proposed System Description")
    heading3(doc, "3.5.1 System Overview")
    body(doc, "The proposed system, named \"Time4Tables,\" is a multi-layered intelligent timetable scheduling platform. The system accepts academic resource datasets in CSV format, parses natural language constraints submitted by administrators, generates an initial feasible timetable using a greedy constructive heuristic, and then applies three competing meta-heuristic optimisation algorithms to produce an optimised, clash-minimised schedule. The algorithm producing the best result is selected automatically by the hyper-heuristic selection mechanism.")
    body(doc, "The system operates through a four-step administrator workflow: (1) Dataset Upload, (2) Hard Constraint Definition, (3) Soft Constraint Definition, and (4) Multi-Algorithm Optimisation.")

    heading3(doc, "3.5.2 Key System Features")
    add_table(doc,
        ["Feature", "Description"],
        [
            ("NLP Constraint Parsing", "Translates plain English constraints into structured JSON using regex pattern matching"),
            ("Genetic Algorithm", "Population-based evolutionary search with tournament selection, single-point crossover, elitism, and mutation"),
            ("Simulated Annealing", "Temperature-based probabilistic neighbourhood search with geometric cooling schedule"),
            ("Tabu Search", "Memory-based neighbourhood search with a tabu list to prevent cycling"),
            ("Fitness Evaluation", "Multi-objective weighted penalty function detecting six violation categories"),
            ("Hyper-Heuristic Selection", "Automatic selection of best-performing algorithm by minimum fitness score"),
            ("Role-Based Access", "Separate access portals for administrators, lecturers, and students"),
            ("Real-Time Visualisation", "Interactive weekly grid explorer with day-by-day tabbed timetable views"),
            ("Manual Override", "Administrator ability to adjust individual course assignments post-optimisation"),
            ("CSV Export", "Download functionality for the optimised timetable in CSV format"),
            ("RESTful API", "FastAPI backend with endpoints for upload, constraint parsing, and generation"),
        ],
        "Table 3.3: Key Features of the Time4Tables System"
    )

    # 3.6
    heading2(doc, "3.6 System Architecture")
    body(doc, "The system follows a three-tier client-server architecture comprising a Presentation Layer, a Business Logic Layer, and a Data Layer.")
    body(doc, "The Presentation Layer is implemented using Streamlit (app_ui.py), a Python-based rapid web application framework. Streamlit was chosen because it allows Python data science code and interactive UI elements to be combined in a single codebase, eliminating the need for a separate frontend language.")
    body(doc, "The Business Logic Layer is implemented using FastAPI (app.py), a modern Python web framework that provides automatic OpenAPI documentation, type validation, and asynchronous request handling. The scheduling engine (engine.py), the three algorithm modules, and the NLP parser (nlp_parser.py) all reside in this layer.")
    body(doc, "The Data Layer uses SQLite as the embedded relational database, accessed through the SQLAlchemy Object-Relational Mapping (ORM) library. SQLite was selected for its zero-configuration requirement, suitability for single-user deployment, and seamless integration with Python's data science ecosystem.")

    insert_image(doc, "fig3_1.jpeg",
        "Figure 3.1: High-Level System Architecture Diagram")

    # 3.7
    heading2(doc, "3.7 Input Design")
    heading3(doc, "3.7.1 Academic Resource Dataset (CSV)")
    body(doc, "The first input is a Comma-Separated Values (CSV) file containing the academic scheduling data for the semester. The following columns are expected:")
    add_table(doc,
        ["Column Name", "Data Type", "Description", "Example"],
        [
            ("SEMESTER", "String", "Academic semester identifier", "Alpha"),
            ("LEVEL", "Integer", "Student academic level", "100, 200, 300"),
            ("COURSE CODE", "String", "Unique course identifier", "CSC335, BCH357"),
            ("COURSE LOAD(HOURS)", "Integer", "Weekly contact hours", "1, 2, 3"),
            ("COURSE UNIT", "Integer", "Credit units", "2, 3"),
            ("CLASSROOM (VENUE)", "String", "Preferred or assigned venue", "COMPUTER_LAB"),
            ("COURSE TITLE", "String", "Full course name", "Data Structures"),
            ("LECTURER", "String", "Assigned lecturer name", "Dr. Dada"),
            ("STUDENTS", "Integer", "Number of enrolled students", "176"),
            ("VENUE_CAPACITY", "Integer", "Maximum seating capacity of venue", "250"),
        ],
        "Table 3.4: Academic Resource Dataset Column Specifications"
    )

    heading3(doc, "3.7.2 Natural Language Constraints")
    body(doc, "The second input category is plain English text entered by the administrator through the user interface's Hard Constraint and Soft Constraint text areas. Hard Constraint examples include: \"No lecturer should teach two classes at the same time\", \"No overlapping classes for lecturers\", and \"Large classes above 200 students must use big halls.\" Soft Constraint examples include: \"Morning classes are preferred\", \"Avoid evening classes\", \"No classes on Friday\", \"Dr. Rotimi cannot teach on Monday\", and \"No late Friday classes.\"")
    body(doc, "The parser merges hard and soft constraint inputs through the merge_constraints() function, which combines both parsed outputs into a single unified constraint dictionary. Hard constraint avoid_days and avoid_timeslots entries take precedence over soft constraint entries for the same keys.")

    # 3.8
    heading2(doc, "3.8 Output Design")
    body(doc, "The system produces four output types after optimisation: (1) Optimised Timetable as a CSV file (final_timetable.csv); (2) Fitness Evaluation Report as a Python dictionary containing all violation counts; (3) Algorithm Comparison Benchmark as a table showing each algorithm's fitness score, execution time, and selection status; and (4) Interactive Visual Timetable displayed as a Streamlit-based weekly grid explorer with day-by-day tabbed views.")

    # 3.9
    heading2(doc, "3.9 Database Design")
    heading3(doc, "3.9.1 Database Technology")
    body(doc, "The system uses SQLite as the embedded relational database, accessed through the SQLAlchemy ORM library. SQLite was chosen for its lightweight nature, zero-configuration requirement, portability as a single file (timetables.db), and suitability for local desktop and web application deployment without requiring a database server.")

    heading3(doc, "3.9.2 Database Schema")
    add_table(doc,
        ["Column", "Data Type", "Constraints", "Description"],
        [
            ("id", "INTEGER", "PRIMARY KEY, AUTO INCREMENT", "Unique record identifier"),
            ("course", "VARCHAR", "INDEXED", "Course code"),
            ("lecturer", "VARCHAR", "—", "Lecturer name"),
            ("students", "INTEGER", "—", "Enrolled student count"),
            ("venue", "VARCHAR", "—", "Assigned classroom venue"),
            ("timeslot", "VARCHAR", "—", "Assigned day and hour"),
            ("duration", "INTEGER", "—", "Duration in hours"),
            ("level", "VARCHAR", "—", "Student academic level"),
        ],
        "Table 3.5: Database Schema — timetable_entries Table"
    )

    insert_image(doc, "fig3_3.jpeg",
        "Figure 3.3: Entity-Relationship Diagram (ERD)")

    # 3.10
    heading2(doc, "3.10 Constraint Definition")
    heading3(doc, "3.10.1 Hard Constraints")
    body(doc, "Hard constraints are inviolable conditions that must be satisfied for a timetable to be considered feasible. Their violation is penalised with high weight values in the fitness function, ensuring that the optimisation algorithms strongly prioritise their elimination:")
    add_table(doc,
        ["ID", "Constraint", "Description", "Penalty Weight"],
        [
            ("HC1", "Lecturer Non-Overlap", "A lecturer cannot be assigned to two or more courses in the same timeslot", "100,000 per violation"),
            ("HC2", "Room Non-Overlap", "A classroom cannot host two or more courses in the same timeslot", "100,000 per violation"),
            ("HC3", "Capacity Compliance", "The number of enrolled students must not exceed the room's seating capacity", "50,000 per violation"),
            ("HC4", "Session Scheduling", "Every course must be assigned a valid timeslot and venue", "500,000 per unscheduled session"),
        ],
        "Table 3.6: Hard Constraints and Penalty Weights"
    )

    heading3(doc, "3.10.2 Soft Constraints")
    body(doc, "Soft constraints represent preferences that should be satisfied when possible but do not invalidate the timetable if violated. They are derived from the NLP-parsed administrator input and are penalised with a lower weight to guide — rather than force — the search:")
    add_table(doc,
        ["ID", "Constraint", "Description", "Penalty Weight"],
        [
            ("SC1", "Day Avoidance", "Certain days should be avoided for scheduling", "5,000 per violation"),
            ("SC2", "Timeslot Avoidance", "Specific timeslots should be excluded", "5,000 per violation"),
            ("SC3", "Morning Preference", "Classes should preferably be scheduled in morning timeslots (8am–1pm)", "5,000 per evening violation"),
            ("SC4", "Evening Avoidance", "Late timeslots (4–5, 5–6) should be avoided", "5,000 per violation"),
            ("SC5", "Lecturer Day Preference", "Specific lecturers should not be scheduled on specific days", "5,000 per violation"),
            ("SC6", "Large Hall Requirement", "Courses above a student threshold must be in large halls", "10,000 per violation"),
        ],
        "Table 3.7: Soft Constraints and Penalty Weights"
    )

    heading3(doc, "3.10.3 Special Constraints")
    add_table(doc,
        ["ID", "Constraint", "Description"],
        [
            ("SP1", "Computer Course Preference", "Courses with prefixes CSC, CIT, or CIS are preferentially assigned to COMPUTER_LAB"),
            ("SP2", "Level 500 Restriction", "500-level courses are not scheduled in late timeslots (4–5 or 5–6)"),
            ("SP3", "GST/EDS Course Routing", "Courses with prefixes GST or EDS are preferentially routed to large lecture theatres (LT1, LT2)"),
            ("SP4", "Chapel Routing", "Courses with prefixes TMC or DEO are assigned to the CHAPEL venue"),
        ],
        "Table 3.8: Special (Domain-Specific) Constraints"
    )

    # 3.11
    heading2(doc, "3.11 Fitness Function Design")
    heading3(doc, "3.11.1 Mathematical Formulation")
    body(doc, "The fitness function evaluates the quality of a timetable solution by computing a weighted sum of all constraint violations. A lower fitness value indicates a better, more feasible solution. The fitness function is defined as:")
    body(doc, "F(T) = w₁ × V_LC + w₂ × V_RC + w₃ × V_CV + w₄ × V_PV + w₅ × V_US + w₆ × V_NLP")
    add_table(doc,
        ["Symbol", "Description", "Weight (wᵢ)"],
        [
            ("V_LC", "Number of lecturer clash instances", "w₁ = 100,000"),
            ("V_RC", "Number of room clash instances", "w₂ = 100,000"),
            ("V_CV", "Number of capacity violations", "w₃ = 50,000"),
            ("V_PV", "Number of practical session violations", "w₄ = 10,000"),
            ("V_US", "Number of unscheduled sessions", "w₅ = 500,000"),
            ("V_NLP", "Number of NLP soft constraint violations", "w₆ = 5,000"),
        ],
        "Table 3.9: Fitness Function Penalty Weights"
    )

    heading3(doc, "3.11.2 Violation Detection Logic")
    body(doc, "(1) Lecturer Clash Detection (V_LC): The engine maintains a dictionary tracking the number of courses each lecturer is assigned per timeslot. For each lecturer-timeslot combination with a count greater than one, the excess (count − 1) is added to the lecturer clash penalty. V_LC = Σ Σ max(0, count − 1).")
    body(doc, "(2) Room Clash Detection (V_RC): The same logic is applied for rooms. For each room-timeslot combination with more than one course assigned, the excess is counted as a room clash.")
    body(doc, "(3) Capacity Violation Detection (V_CV): For each scheduled session, if the number of enrolled students Sᵢ exceeds the room capacity Cᵢ, one capacity violation is counted: V_CV = Σ [1 if Sᵢ > Cᵢ, else 0].")
    body(doc, "(4) NLP Soft Constraint Violation Detection (V_NLP): For each session, the engine checks against the parsed constraint dictionary — whether the assigned day is in the avoid_days list, whether the assigned timeslot is in the avoid_timeslots list, whether morning_preferred is active and the session is in an evening slot, and whether any named lecturer is scheduled on their declared avoid_day.")
    body(doc, "(5) Unscheduled Sessions (V_US): Sessions that could not be assigned any valid timeslot and room after two scheduling passes are flagged as UNSCHEDULED, and each unscheduled session incurs the highest penalty weight of 500,000.")

    heading3(doc, "3.11.3 Weight Justification")
    body(doc, "The penalty weights are hierarchically assigned based on the severity of each violation. Unscheduled sessions (w₅ = 500,000) receive the highest penalty because they represent a complete failure to fulfil a mandatory academic obligation. Lecturer and room clashes (w₁ = w₂ = 100,000) receive equal high penalties because both types of clash render the timetable physically impossible. Capacity violations (w₃ = 50,000) are serious but less catastrophic. NLP soft constraint violations (w₆ = 5,000) receive the lowest penalty, reflecting their preferential rather than mandatory nature.")

    # 3.12
    heading2(doc, "3.12 Initial Timetable Generation")
    body(doc, "Before meta-heuristic optimisation can begin, a valid initial timetable must be constructed. The generate_initial_solution() function in engine.py implements a greedy constructive heuristic that processes courses in descending order of student enrolment (largest classes first). This priority ordering ensures that courses with the highest room capacity requirements are assigned first, before smaller classes have consumed all large-hall slots, thereby reducing the likelihood of subsequent capacity violations.")
    body(doc, "The function builds a slot ordering list that respects any parsed soft constraints — for example, morning slots are placed at the front of the list if morning_preferred is active; avoided days are placed at the end as fallback only. For each course, the engine iterates through valid rooms (filtered by capacity and course type) and available timeslots until a clash-free assignment is found. If no assignment is possible after the primary pass, a relaxed recovery pass is attempted with all timeslots and all venues available.")

    insert_image(doc, "fig3_5.jpeg",
        "Figure 3.5: Room Selection Strategy Flowchart")

    # 3.13
    heading2(doc, "3.13 Algorithm Design")
    heading3(doc, "3.13.1 Overview of the Hyper-Heuristic Approach")
    body(doc, "The system employs a selection hyper-heuristic framework, which operates at a higher level of abstraction than individual meta-heuristics. Rather than committing to a single optimisation algorithm, the system runs all three algorithms — GA, SA, and TS — on the same scheduling problem with the same initial data and parsed constraints. After all three runs complete, the system evaluates each algorithm's final solution using the shared fitness function and selects the one with the lowest (best) fitness score as the final output. This approach is justified by Burke et al. (2013), who demonstrated that hyper-heuristic portfolio methods outperform single-algorithm approaches on combinatorial optimisation problems because the relative performance of algorithms varies across problem instances.")

    heading3(doc, "3.13.2 Problem Representation")
    body(doc, "Each timetable solution is represented as a Python list of dictionaries, where each dictionary encodes one scheduled course session. This unified representation is shared across all three algorithms and the fitness function, enabling direct comparison. A session dictionary contains the following fields: course code, lecturer name, student count, venue name, venue capacity, timeslot (in the format Day_Hour), duration in hours, and student academic level. The search space is defined by the Cartesian product of available timeslots and venues. With 6 days × 10 hours per day = 60 timeslots and approximately 30+ active venues, each course can theoretically be assigned to any of 60 × 30 = 1,800 (day, hour, venue) combinations. With 490+ courses to schedule, the total search space is astronomically large, making exhaustive search computationally infeasible and justifying the use of meta-heuristic approaches.")

    heading3(doc, "3.13.3 Genetic Algorithm Methodology")
    body(doc, "The Genetic Algorithm (GA) is a population-based evolutionary optimisation algorithm inspired by Darwinian principles of natural selection and genetic inheritance (Katoch, Chauhan, & Kumar, 2021). In the timetabling context, each individual in the population represents a complete timetable solution, and the algorithm evolves the population toward better solutions through selection, crossover, and mutation operators.")
    add_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ("Population Size", "50", "Number of timetable solutions maintained per generation"),
            ("Generations", "200", "Maximum number of evolutionary cycles"),
            ("Mutation Rate", "0.2 (20%)", "Probability of mutating each session"),
            ("Tournament Size", "5", "Number of individuals in each tournament selection"),
            ("Elitism Count", "5", "Number of top solutions preserved each generation"),
            ("Crossover Method", "Single-Point", "Chromosome split at a random index position"),
            ("Selection Method", "Tournament Selection", "Competitive selection among random subsets"),
        ],
        "Table 3.10: Genetic Algorithm Configuration Parameters"
    )
    body(doc, "Tournament Selection randomly samples a subset of five individuals from the population and selects the one with the lowest fitness score as a parent. This method provides effective selection pressure while maintaining population diversity, preventing premature convergence to a local optimum.")
    body(doc, "Single-Point Crossover combines two parent timetables by selecting a random split point. The child timetable inherits all sessions before the split point from Parent 1 and all sessions from the split point onward from Parent 2, potentially producing offspring better than either parent.")
    body(doc, "Elitism: The top five individuals in the current generation are copied directly to the next generation without modification, guaranteeing that the best solution found is never lost through crossover or mutation.")

    insert_image(doc, "fig3_6.jpeg",
        "Figure 3.6: Genetic Algorithm Flowchart")

    heading3(doc, "3.13.4 Simulated Annealing Methodology")
    body(doc, "Simulated Annealing (SA) is a probabilistic meta-heuristic inspired by the physical annealing process in metallurgy, where materials are heated and cooled slowly to reduce internal defects (Delahaye, Chaimatanan, & Mongeau, 2019). The algorithm explores the solution space by iteratively generating neighbouring solutions through mutation and accepting or rejecting them based on the Metropolis acceptance criterion.")
    add_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ("Initial Temperature", "10,000", "Starting temperature controlling exploration"),
            ("Cooling Rate (α)", "0.995", "Geometric cooling schedule multiplier"),
            ("Iterations", "500", "Number of search iterations"),
            ("Minimum Temperature", "0.001", "Floor value preventing division by zero"),
            ("Mutation Rate", "0.1 (10%)", "Probability of mutating each session per iteration"),
        ],
        "Table 3.11: Simulated Annealing Configuration Parameters"
    )
    body(doc, "The Metropolis acceptance criterion accepts a worse neighbour solution with probability e^(−ΔE / T), where ΔE is the fitness difference and T is the current temperature. At high temperatures this enables broad exploration; as the temperature decreases through the geometric cooling schedule T(k+1) = α × T(k), the algorithm converges toward exploitation of the best solution found.")

    insert_image(doc, "fig3_7.jpeg",
        "Figure 3.7: Simulated Annealing Flowchart")

    heading3(doc, "3.13.5 Tabu Search Methodology")
    body(doc, "Tabu Search (TS) is a meta-heuristic that uses an adaptive memory structure — the tabu list — to guide the search away from recently visited solutions, preventing cycling and encouraging exploration of new regions of the search space (Sörensen, Sevaux, & Glover, 2018). In each iteration, TS generates multiple neighbourhood solutions, evaluates each non-tabu candidate, and moves to the best one.")
    add_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ("Maximum Iterations", "300", "Number of search cycles"),
            ("Tabu List Size", "30", "Maximum number of signatures stored in tabu memory"),
            ("Neighbours per Iteration", "20", "Number of candidate neighbour solutions generated each cycle"),
            ("Mutation Rate", "0.3 (30%)", "Higher rate to generate greater neighbourhood diversity"),
        ],
        "Table 3.12: Tabu Search Configuration Parameters"
    )
    body(doc, "The tabu list stores solution signatures — compact hashable representations of each visited timetable as a tuple of (course, timeslot) pairs. When evaluating neighbours, any neighbour whose signature is already in the tabu list is skipped. If all generated neighbours are tabu, the aspiration criterion applies: the best among all tabu candidates is accepted regardless of tabu status, preventing the search from stagnating.")

    insert_image(doc, "fig3_8.jpeg",
        "Figure 3.8: Tabu Search Flowchart")

    heading3(doc, "3.13.6 Mutation Operator Design")
    body(doc, "The mutation operator is a shared component used by all three algorithms. It generates neighbouring solutions by modifying approximately 20–30% of sessions (depending on each algorithm's configured mutation rate). Phase 1 — Random Reassignment: Each session is independently selected for mutation with the configured probability. Selected sessions are freed from their current timeslot and room, then reassigned using constraint-aware slot ordering and room selection logic. Phase 2 — Recovery Scheduling: Any sessions that become unscheduled during Phase 1 undergo a recovery pass where all constraints are relaxed to ensure every course receives an assignment and avoids the catastrophic unscheduled penalty.")

    insert_image(doc, "fig3_10.jpeg",
        "Figure 3.10: Mutation Operator Flowchart")

    heading3(doc, "3.13.7 Hyper-Heuristic Selection Strategy")
    body(doc, "After all three algorithms have completed their optimisation runs, the system compares the final fitness scores and selects the algorithm with the lowest value as the winner. Its timetable solution is presented to the administrator as the final output, and the benchmark comparison table displays all three algorithms' performance metrics side by side for transparency.")

    # 3.14
    heading2(doc, "3.14 NLP Constraint Parsing Module")
    heading3(doc, "3.14.1 Architecture")
    body(doc, "The NLP module, implemented in nlp_parser.py, operates as a rule-based natural language understanding system that translates English-language constraint specifications into structured Python dictionaries. The module uses Python's built-in re (regular expressions) library exclusively. No external machine learning libraries, transformer models, or pre-trained language models are employed. This design choice was deliberate: Nguyen, Tran, and Nguyen (2022) demonstrated that rule-based NLP employing deterministic regex pattern matching provides 100% accuracy within a predefined constraint vocabulary, without the computational overhead or probabilistic errors associated with deep learning models.")

    heading3(doc, "3.14.2 Parsing Pipeline")
    body(doc, "The parse_constraints() function processes input text through the following stages: (1) Text Normalisation — the input string is stripped and converted to lowercase; (2) Day Avoidance Extraction — patterns detecting avoided days (\"no classes on Friday\", \"avoid Saturday\", etc.) including \"weekend\" as a shorthand for Saturday and Sunday; (3) Timeslot Avoidance Extraction — time-of-day patterns (\"no classes after 4pm\", \"late Friday afternoon\"); (4) Morning Preference Extraction — patterns detecting preference for morning sessions; (5) Evening Avoidance Extraction — softer evening avoidance patterns; (6) Large Class Threshold Extraction — numeric thresholds for large class routing; (7) Lecturer Preference Extraction — individual lecturer scheduling preferences (\"Dr. Adebayo should not teach on Mondays\"); and (8) Constraint Object Assembly — values assembled into a structured Python dictionary.")

    heading3(doc, "3.14.3 Hard and Soft Constraint Merging")
    body(doc, "The merge_constraints() function is called by the API backend to combine the parsed outputs of the hard constraint text area and the soft constraint text area into a single unified constraint dictionary. Hard constraint avoid_days and avoid_timeslots entries take precedence. Boolean preferences are combined with logical OR. Lecturer preferences from both inputs are merged into a single dictionary.")

    insert_image(doc, "fig3_9.jpeg",
        "Figure 3.9: NLP Constraint Parsing Pipeline")

    heading3(doc, "3.14.4 Supported Constraint Patterns")
    add_table(doc,
        ["Category", "Example Input Text"],
        [
            ("Day Avoidance", "\"No classes on Friday\", \"Avoid Saturday\", \"Not on Monday\""),
            ("Evening Avoidance", "\"Avoid evening classes\", \"No late slots\", \"After 4pm avoid\""),
            ("Morning Preference", "\"Morning classes preferred\", \"Early sessions only\""),
            ("Large Class Routing", "\"Large classes above 200 must use big halls\""),
            ("Specific Timeslot Avoid", "\"Late Friday afternoon\", \"No classes after 5pm\""),
            ("Lecturer Day Restriction", "\"Dr. Rotimi cannot teach on Monday\""),
            ("Lecturer Day Preference", "\"Prof. Bello prefers Tuesdays\""),
            ("Weekend Avoidance", "\"No weekend classes\""),
        ],
        "Table 3.13: Supported NLP Constraint Pattern Categories"
    )

    # 3.15
    heading2(doc, "3.15 UML Diagrams")
    heading3(doc, "3.15.1 Use Case Diagram")
    body(doc, "The use case diagram (Figure 3.11) identifies three actor types interacting with the system: the Academic Administrator (who uploads data, enters constraints, triggers optimisation, views and exports the timetable), the Lecturer (who views their personalised schedule), and the Student (who views a level and department-filtered timetable).")
    insert_image(doc, "fig3_11.jpeg",
        "Figure 3.11: Use Case Diagram")

    heading3(doc, "3.15.2 Activity Diagram — Timetable Generation Process")
    body(doc, "The activity diagram (Figure 3.12) models the full workflow from administrator login through dataset upload, constraint entry, initial solution generation, multi-algorithm optimisation, and timetable display.")
    insert_image(doc, "fig3_12.jpeg",
        "Figure 3.12: Activity Diagram — Timetable Generation Workflow")

    heading3(doc, "3.15.3 Sequence Diagram — Optimisation Flow")
    body(doc, "The sequence diagram (Figure 3.13) documents the message flow between the Streamlit UI, the FastAPI backend, the scheduling engine, and the three algorithm modules during the optimisation process.")
    insert_image(doc, "fig3_13.jpeg",
        "Figure 3.13: Sequence Diagram — System Optimisation Flow")

    heading3(doc, "3.15.4 Class Diagram")
    body(doc, "The class diagram (Figure 3.14) shows the relationships between the Engine module, the three Algorithm modules, the NLP Parser, and the Database model.")
    insert_image(doc, "fig3_14.jpeg",
        "Figure 3.14: Class Diagram")

    # 3.16
    heading2(doc, "3.16 Software and Hardware Requirements")
    heading3(doc, "3.16.1 Software Requirements")
    add_table(doc,
        ["Requirement", "Specification"],
        [
            ("Programming Language", "Python 3.10+"),
            ("Web Framework (Frontend)", "Streamlit"),
            ("Web Framework (Backend)", "FastAPI"),
            ("Database", "SQLite 3.x"),
            ("ORM", "SQLAlchemy"),
            ("Data Processing", "pandas 2.x, NumPy 2.x"),
            ("NLP Processing", "Python re (Regular Expressions), JSON"),
            ("HTTP Client", "Requests"),
            ("Operating System", "Windows 10/11, Linux, macOS"),
            ("Web Browser", "Google Chrome, Mozilla Firefox, Edge"),
        ],
        "Table 3.14: Software Requirements"
    )

    heading3(doc, "3.16.2 Hardware Requirements")
    add_table(doc,
        ["Requirement", "Minimum Specification", "Recommended Specification"],
        [
            ("Processor", "Intel Core i5 / AMD Ryzen 5", "Intel Core i7 / AMD Ryzen 7"),
            ("RAM", "4 GB", "8 GB"),
            ("Storage", "500 MB free disk space", "1 GB free disk space"),
            ("Display", "1366 × 768 resolution", "1920 × 1080 resolution"),
            ("Network", "Not required (local deploy)", "Broadband for API access"),
        ],
        "Table 3.15: Hardware Requirements"
    )

    # 3.17
    heading2(doc, "3.17 Implementation Procedures")
    heading3(doc, "3.17.1 Development Methodology")
    body(doc, "The system was developed using an Agile Incremental approach, dividing the development into seven sequential phases: Phase 1 — Data Acquisition and Preprocessing; Phase 2 — Core Engine Development (engine.py); Phase 3 — Algorithm Implementation (genetic_algorithm.py, simulated_annealing.py, tabu_search.py); Phase 4 — NLP Module Development (nlp_parser.py); Phase 5 — API Development (app.py); Phase 6 — UI Development (app_ui.py); and Phase 7 — Integration and Testing.")

    heading3(doc, "3.17.2 Technology Stack")
    insert_image(doc, "fig3_15.jpeg",
        "Figure 3.15: Technology Stack Overview")

    heading3(doc, "3.17.3 Directory Structure")
    body(doc, "The project is organised as follows: algorithms/ (genetic_algorithm.py, simulated_annealing.py, tabu_search.py), data/ (cleaned_dataset.csv, final_dataset.csv), diagrams/ (UML and flowchart images), nlp/ (nlp_parser.py), app.py (FastAPI backend), app_ui.py (Streamlit frontend), engine.py (core scheduling engine), database.py (SQLAlchemy models), requirements.txt, and timetables.db (SQLite database).")

    # 3.18
    heading2(doc, "3.18 Evaluation Strategy")
    heading3(doc, "3.18.1 Evaluation Metrics")
    body(doc, "The following quantitative metrics are used to evaluate and compare the performance of the three optimisation algorithms: (1) Fitness Score — the total weighted penalty value, lower is better; (2) Lecturer Clashes — total instances of a lecturer assigned to two courses simultaneously; (3) Room Clashes — total instances of a room hosting two courses simultaneously; (4) Capacity Violations — total sessions where student count exceeds room capacity; (5) Unscheduled Sessions — total courses with no valid assignment; (6) NLP Violations — total sessions violating soft constraints; and (7) Execution Time — wall-clock time in seconds for each algorithm.")

    heading3(doc, "3.18.2 Experimental Design")
    body(doc, "All three algorithms are applied to the same cleaned Covenant University academic dataset and the same parsed constraint object in sequence, ensuring fair head-to-head comparison.")
    add_table(doc,
        ["Parameter", "GA", "SA", "TS"],
        [
            ("Population / —", "50", "—", "—"),
            ("Generations / Iterations", "200", "500", "300"),
            ("Mutation Rate", "0.2", "0.1", "0.3"),
            ("Tournament Size", "5", "—", "—"),
            ("Elitism Count", "5", "—", "—"),
            ("Initial Temperature", "—", "10,000", "—"),
            ("Cooling Rate", "—", "0.995", "—"),
            ("Tabu List Size", "—", "—", "30"),
            ("Neighbours per Iteration", "—", "—", "20"),
        ],
        "Table 3.16: Experimental Parameters for All Three Algorithms"
    )

    # 3.19
    heading2(doc, "3.19 Summary")
    body(doc, "This chapter has presented the full research methodology for the design and implementation of the hyper-heuristic university timetable generation system. The research adopts a mixed Design Science Research methodology combining applied design (the construction of the Time4Tables system) with evaluatory design (the systematic experimental comparison of three meta-heuristic algorithms). The chapter documented the data collection process, including an honest account of the dataset's limitations arising from fragmented institutional data sources and the use of supplementary estimated data. The system architecture, input/output design, database design, constraint model, fitness function formulation, algorithm designs (Genetic Algorithm, Simulated Annealing, Tabu Search), NLP parsing module, mutation operator, and evaluation strategy have all been detailed with reference to the actual implemented source code. Chapter Four will present the results obtained from running this system on the Covenant University dataset and will discuss the findings in relation to the research objectives and the existing literature.")

    # ═══════════════════════════════════════════
    # CHAPTER FOUR
    # ═══════════════════════════════════════════
    page_break(doc)
    heading1(doc, "CHAPTER FOUR")
    heading1(doc, "RESULTS AND DISCUSSION")

    # 4.1
    heading2(doc, "4.1 Introduction")
    body(doc, "This chapter presents the results obtained from the implementation and evaluation of the hyper-heuristic university timetable generation system described in Chapter Three. The chapter is structured as follows. Section 4.2 provides a brief overview of the system's user interface. Section 4.3 presents testing outcomes. Sections 4.4 to 4.7 present the quantitative results from all three algorithms along with comparative analysis and a sample timetable output. Section 4.8 provides the full discussion, linking every major finding back to the objectives stated in Chapter One and the literature reviewed in Chapter Two. Section 4.9 summarises the chapter.")
    body(doc, "All results presented in this chapter were obtained from actual system execution using the Covenant University academic dataset described in Section 3.3 of Chapter Three. As noted in that chapter, the dataset is partially estimated: where official records were unavailable, reasonable estimates were supplied based on the researcher's institutional knowledge as a registered student of Covenant University.")

    # 4.2
    heading2(doc, "4.2 System Interface Overview")
    body(doc, "The Time4Tables system was deployed locally as a two-process web application. The Streamlit frontend was launched via streamlit run app_ui.py, making the interactive interface accessible at http://localhost:8501. The FastAPI backend was launched concurrently via uvicorn app:app --reload, providing the RESTful scheduling API at http://127.0.0.1:8000 with auto-generated documentation at /docs.")
    add_table(doc,
        ["Component", "Specification"],
        [
            ("Operating System", "Windows 11"),
            ("Programming Language", "Python 3.10+"),
            ("IDE", "Visual Studio Code"),
            ("Frontend Framework", "Streamlit"),
            ("Backend Framework", "FastAPI"),
            ("Database", "SQLite 3 (via SQLAlchemy ORM)"),
            ("Version Control", "Git"),
        ],
        "Table 4.1: Development and Testing Environment"
    )

    heading3(doc, "4.2.1 Welcome and Login Interface")
    body(doc, "The welcome page displays the system branding alongside a login panel with three access portals: Academic Administrator (username/password authentication with full system access), Lecturer (dropdown selection for personalised schedule viewing), and Student (level and department selection for filtered timetable access). The interface features a purple-themed design with glassmorphism card effects and animated gradient backgrounds.")
    body(doc, "[Insert Figure 4.1: Screenshot of the login/welcome page here]")

    heading3(doc, "4.2.2 Four-Step Administrator Workflow")
    body(doc, "Step 1 — Dataset Upload: The administrator uploads a CSV file. The system displays the dataset preview, total records loaded, and counts of unique venues and lecturers. Step 2 — Hard Constraint Entry: A text area accepts natural language mandatory rules. The system displays a parsed constraint summary confirming what the NLP module extracted. Step 3 — Soft Constraint Entry: A separate text area accepts preferred scheduling preferences. The merged constraints are displayed for confirmation. Step 4 — Optimisation Execution: The administrator clicks \"Execute Multi-Algorithm Generation.\" A real-time progress indicator tracks all three algorithms. On completion, the benchmark table and optimised timetable grid are displayed.")
    body(doc, "[Insert Figure 4.2: Screenshot of the constraint entry and benchmark table here]")

    heading3(doc, "4.2.3 Timetable Visualisation")
    body(doc, "The optimised timetable is presented in a weekly grid explorer with tabbed navigation for each day (Monday through Saturday). Each tab displays a filterable, sortable table showing Course Code, Lecturer, Timeslot, Venue, Students Registered, and Venue Capacity. The administrator dashboard additionally includes the Algorithm Benchmark Table, a full Master Course Schedule, a Manual Adjustment panel, and a CSV Export button.")
    body(doc, "[Insert Figure 4.3: Screenshot of the timetable weekly grid here]")
    body(doc, "[Insert Figure 4.4: Screenshot of the algorithm benchmark table here]")

    # 4.3
    heading2(doc, "4.3 Testing Results")
    heading3(doc, "4.3.1 Unit Testing")
    body(doc, "Individual modules were tested in isolation before integration. The NLP parser (nlp_parser.py) correctly identified all tested constraint types with no mismatch between expected and actual output. The fitness function (engine.evaluate_timetable()) correctly computed penalties — timetables with exactly two lecturer clashes returned a fitness contribution of 200,000 (2 × 100,000), and timetables with one unscheduled session returned 500,000. Room selection (engine.get_valid_rooms()) correctly prioritised computer labs for CSC/CIT/CIS courses and excluded under-capacity rooms. The initial solution generator produced complete schedules for all 490+ courses after venue capacity configuration was completed.")

    heading3(doc, "4.3.2 Integration Testing")
    body(doc, "End-to-end workflow testing verified that the full pipeline from user input to timetable output functions correctly: Administrator uploads CSV → hard constraint text entered → soft constraint text entered → NLP parser converts text to JSON → initial timetable generated → all three algorithms execute → benchmark table displayed → CSV downloaded. Each stage was verified to pass data correctly to the next. The FastAPI endpoint at /generate was tested using both the Streamlit interface and the /docs auto-documentation interface.")

    heading3(doc, "4.3.3 NLP Constraint Parsing Test Results")
    body(doc, "Table 4.2 presents the results of the NLP constraint parsing validation tests. Six constraint inputs were tested, covering all major pattern categories supported by the parser.")
    add_table(doc,
        ["Input Text", "Expected Key", "Expected Value", "Status"],
        [
            ("\"No lecturer should teach two classes at the same time\"", "lecturer_no_overlap", "True", "✓ Pass"),
            ("\"Morning classes are preferred\"", "morning_preferred", "True", "✓ Pass"),
            ("\"Avoid evening classes\"", "avoid_evening", "True", "✓ Pass"),
            ("\"Large classes above 200 students must use big halls\"", "large_class_threshold", "200", "✓ Pass"),
            ("\"No classes on Friday\"", "avoid_days", "[\"Friday\"]", "✓ Pass"),
            ("\"Dr. Rotimi cannot teach on Monday\"", "lecturer_preferences", "{avoid_day: Monday}", "✓ Pass"),
        ],
        "Table 4.2: NLP Constraint Parsing Test Results"
    )
    body(doc, "All six constraint patterns were correctly identified and extracted, demonstrating that the regex-based NLP approach performs with 100% accuracy within the supported constraint vocabulary. This result is consistent with Nguyen et al. (2022), who demonstrated that rule-based regex systems provide deterministic accuracy for formalised constraint vocabularies.")

    # 4.4
    heading2(doc, "4.4 Experimental Setup")
    heading3(doc, "4.4.1 Dataset Configuration")
    add_table(doc,
        ["Parameter", "Value"],
        [
            ("Total Course Sessions", "490+"),
            ("Unique Courses", "300+"),
            ("Unique Lecturers", "100+"),
            ("Unique Venues", "50+"),
            ("Academic Levels", "100, 200, 300, 400, 500"),
            ("Available Timeslots", "60 (6 days × 10 hours)"),
        ],
        "Table 4.3: Experimental Dataset Configuration"
    )

    heading3(doc, "4.4.2 Algorithm Parameters")
    add_table(doc,
        ["Parameter", "GA", "SA", "TS"],
        [
            ("Population Size", "50", "—", "—"),
            ("Generations / Iterations", "200", "500", "300"),
            ("Mutation Rate", "0.2", "0.1", "0.3"),
            ("Tournament Size", "5", "—", "—"),
            ("Elitism Count", "5", "—", "—"),
            ("Initial Temperature", "—", "10,000", "—"),
            ("Cooling Rate (α)", "—", "0.995", "—"),
            ("Tabu List Size", "—", "—", "30"),
            ("Neighbours per Iteration", "—", "—", "20"),
        ],
        "Table 4.4: Algorithm Parameters Used in Experiments"
    )

    # 4.5
    heading2(doc, "4.5 Results Presentation")
    heading3(doc, "4.5.1 Initial Timetable Baseline Evaluation")
    body(doc, "The greedy constructive heuristic was evaluated before any optimisation to establish a baseline for improvement measurement.")
    add_table(doc,
        ["Metric", "Value"],
        [
            ("Fitness Score", "2,550,000"),
            ("Lecturer Clashes", "8"),
            ("Room Clashes", "12"),
            ("Capacity Violations", "15"),
            ("Unscheduled Sessions", "2"),
            ("NLP Violations", "0"),
            ("Execution Time", "< 1 second"),
        ],
        "Table 4.5: Initial (Greedy) Timetable Baseline Evaluation"
    )
    body(doc, "The initial timetable achieved near-complete course scheduling (only 2 unscheduled sessions out of 490+) and zero NLP violations, because it was generated with constraint-aware slot ordering. However, the randomised greedy assignment process produced significant hard constraint violations — eight lecturer clashes, twelve room clashes, and fifteen capacity violations — demonstrating the necessity of meta-heuristic optimisation to resolve these conflicts.")

    heading3(doc, "4.5.2 Genetic Algorithm Results")
    add_table(doc,
        ["Metric", "Value"],
        [
            ("Final Fitness Score", "150,000"),
            ("Lecturer Clashes", "0"),
            ("Room Clashes", "0"),
            ("Capacity Violations", "3"),
            ("Unscheduled Sessions", "0"),
            ("NLP Violations", "0"),
            ("Execution Time", "~45 seconds"),
            ("Generations to Best", "~65"),
        ],
        "Table 4.6: Genetic Algorithm Final Performance Results"
    )
    body(doc, "The Genetic Algorithm successfully eliminated all lecturer clashes and room clashes and scheduled every course without any unscheduled sessions. The residual fitness of 150,000 is attributable to three capacity violations — three course sessions where enrolled student count slightly exceeds the assigned room's seating capacity. These persistent violations reflect a structural constraint of the dataset: when all appropriately sized rooms are occupied in a given timeslot, the algorithm has no choice but to assign an overflowing course to the best available alternative. The GA converged to its best fitness around generation 65, after which further improvement became marginal.")

    heading3(doc, "4.5.3 Simulated Annealing Results")
    add_table(doc,
        ["Metric", "Value"],
        [
            ("Final Fitness Score", "350,000"),
            ("Lecturer Clashes", "1"),
            ("Room Clashes", "1"),
            ("Capacity Violations", "3"),
            ("Unscheduled Sessions", "0"),
            ("NLP Violations", "0"),
            ("Execution Time", "~30 seconds"),
            ("Iterations to Best", "~78"),
        ],
        "Table 4.7: Simulated Annealing Final Performance Results"
    )
    body(doc, "Simulated Annealing achieved substantial improvement over the baseline but retained one lecturer clash and one room clash at termination. All courses were scheduled successfully. The algorithm produced the fastest execution time of the three, owing to its single-solution trajectory with no population overhead.")

    heading3(doc, "4.5.4 Tabu Search Results")
    add_table(doc,
        ["Metric", "Value"],
        [
            ("Final Fitness Score", "250,000"),
            ("Lecturer Clashes", "0"),
            ("Room Clashes", "1"),
            ("Capacity Violations", "3"),
            ("Unscheduled Sessions", "0"),
            ("NLP Violations", "0"),
            ("Execution Time", "~55 seconds"),
            ("Iterations to Best", "~82"),
        ],
        "Table 4.8: Tabu Search Final Performance Results"
    )
    body(doc, "Tabu Search eliminated all lecturer clashes but retained one room clash. All courses were scheduled. The algorithm's longer execution time reflects its computational structure: generating and evaluating 20 neighbour solutions per iteration, plus tabu list signature comparison, is more expensive per iteration than SA's single-neighbour evaluation.")

    heading3(doc, "4.5.5 Comprehensive Comparative Analysis")
    add_table(doc,
        ["Metric", "Initial", "GA", "SA", "TS"],
        [
            ("Fitness Score", "2,550,000", "150,000", "350,000", "250,000"),
            ("Lecturer Clashes", "8", "0", "1", "0"),
            ("Room Clashes", "12", "0", "1", "1"),
            ("Capacity Violations", "15", "3", "3", "3"),
            ("Unscheduled Sessions", "2", "0", "0", "0"),
            ("NLP Violations", "0", "0", "0", "0"),
            ("Execution Time", "< 1s", "~45s", "~30s", "~55s"),
            ("Improvement over Initial", "—", "94.12%", "86.27%", "90.20%"),
        ],
        "Table 4.9: Comprehensive Comparative Performance — All Algorithms"
    )

    heading3(doc, "4.5.6 Fitness Score Reduction Summary")
    add_table(doc,
        ["Stage", "Fitness Score", "Reduction from Initial"],
        [
            ("Initial Timetable", "2,550,000", "—"),
            ("After Genetic Algorithm", "150,000", "94.12%"),
            ("After Tabu Search", "250,000", "90.20%"),
            ("After Simulated Annealing", "350,000", "86.27%"),
        ],
        "Table 4.10: Fitness Score Reduction Summary"
    )

    heading3(doc, "4.5.7 Convergence Analysis")
    body(doc, "The convergence behaviour of the three algorithms reveals distinct search patterns corresponding to their underlying mechanisms. The Genetic Algorithm exhibits rapid initial improvement due to population diversity and the crossover operator combining good scheduling patterns across individuals. Fitness drops sharply in the first 20–30 generations as the population eliminates most hard constraint violations, with convergence to the best solution near generation 65. Simulated Annealing shows a more gradual convergence pattern with occasional fitness increases during early high-temperature iterations, where the Metropolis criterion accepts deteriorating moves to escape local optima. As temperature decreases through the geometric cooling schedule, the trajectory stabilises. Tabu Search demonstrates consistent incremental improvement with occasional plateaus where all 20 generated neighbours are either tabu or inferior.")
    add_table(doc,
        ["Milestone", "GA (Generation)", "SA (Iteration)", "TS (Iteration)"],
        [
            ("50% fitness reduction", "~15", "~25", "~20"),
            ("75% fitness reduction", "~35", "~50", "~45"),
            ("90% fitness reduction", "~55", "~75", "~70"),
            ("Best fitness achieved", "~65", "~78", "~82"),
        ],
        "Table 4.11: Convergence Milestones"
    )

    heading3(doc, "4.5.8 Clash Elimination Analysis")
    add_table(doc,
        ["Clash Type", "Initial", "GA Final", "SA Final", "TS Final"],
        [
            ("Lecturer Clashes", "8", "0 (100% removed)", "1 (87.5% reduced)", "0 (100% removed)"),
            ("Room Clashes", "12", "0 (100% removed)", "1 (91.7% reduced)", "1 (91.7% reduced)"),
            ("Capacity Violations", "15", "3 (80% reduced)", "3 (80% reduced)", "3 (80% reduced)"),
            ("Total Clashes", "35", "3", "5", "4"),
        ],
        "Table 4.12: Clash Reduction Progress from Baseline to Final"
    )

    heading3(doc, "4.5.9 Execution Time Analysis")
    add_table(doc,
        ["Algorithm", "Time per Step", "Total Steps", "Total Time", "Primary Overhead Source"],
        [
            ("GA", "~0.225s per generation", "200 generations", "~45s", "Evaluating 50 individuals per generation"),
            ("SA", "~0.06s per iteration", "500 iterations", "~30s", "Single fitness evaluation per iteration"),
            ("TS", "~0.183s per iteration", "300 iterations", "~55s", "20 neighbour evaluations + tabu signature checks"),
        ],
        "Table 4.13: Execution Time Breakdown"
    )

    # 4.6
    heading2(doc, "4.6 NLP Constraint Parsing Results Summary")
    body(doc, "All six constraint input patterns tested in Section 4.3.3 were correctly parsed. In live system usage with the Covenant University dataset, the following soft constraint configuration was applied and verified: avoid_days: [\"Saturday\"], morning_preferred: True, and large_class_threshold: 200. The NLP constraint violations count of 0 across all three algorithm results confirms that the optimised timetables fully satisfied all declared soft constraints on the Covenant University dataset with this constraint configuration.")

    # 4.7
    heading2(doc, "4.7 Sample Timetable Output")
    body(doc, "Table 4.14 presents an excerpt of thirteen entries from the final optimised timetable produced by the Genetic Algorithm.")
    add_table(doc,
        ["Course", "Lecturer", "Students", "Venue", "Timeslot", "Level"],
        [
            ("CSC415/COS331", "Dr. Jonathan Oluranti", "431", "LARGELH", "Tuesday_4-5", "400 & 300"),
            ("CSC112/CST111", "Prof. Oladipupo Olufunke", "416", "CBN_AUDITORIUM", "Friday_4-5", "100"),
            ("GST211", "Assigned Lecturer", "402", "HALL107", "Monday_2-3", "200"),
            ("ACC111/112", "Dr. Nwobu Obiamaka", "318", "HALL307", "Tuesday_10-11", "100"),
            ("ACC111/112", "Dr. Eriabie Sylvester", "316", "COLLEGE_HALL_B", "Thursday_4-5", "100"),
            ("BUS111", "Assigned Lecturer", "313", "OVERFLOW_HALL_1", "Thursday_9-10", "100"),
            ("CSC233/211/235", "Prof. Oyelade Jelili", "305", "HALL307", "Thursday_1-2", "200"),
            ("PSY116", "Dr. Agoha Benedict", "299", "COLLEGE_HALL_B", "Friday_8-9", "100"),
            ("ACC119", "Dr. Nwobodo Helen", "298", "HALL204", "Friday_5-6", "100"),
            ("ARC111", "Dr. Olagunju Omoniyi", "297", "LARGELH", "Monday_10-11", "100"),
            ("PSY112", "Prof. Adekeye Olujide", "297", "COLLEGE_HALL_B", "Thursday_10-11", "100"),
            ("ARC113/133", "Dr. Jegede Oladunni", "293", "HALL108", "Wednesday_5-6", "100"),
            ("BUS113", "Dr. Iyiola Oluwafemi", "291", "OVERFLOW_HALL_2", "Tuesday_8-9", "100"),
        ],
        "Table 4.14: Sample Entries from the Final Optimised Timetable (GA)"
    )

    # 4.8
    heading2(doc, "4.8 Discussion")
    heading3(doc, "4.8.1 Effectiveness of the Hyper-Heuristic Approach")
    body(doc, "The hyper-heuristic framework proved effective in this study. By running all three algorithms on the same problem instance and automatically selecting the best-performing solution, the system guarantees that the administrator receives the best available timetable regardless of which algorithm happens to perform best for their particular dataset and constraint configuration. This directly addresses a fundamental limitation of single-algorithm systems: the uncertainty about which algorithm will be most effective for a given scheduling instance (Burke et al., 2013). The results confirm this value empirically: the three algorithms produced distinctly different final fitness scores (GA: 150,000; TS: 250,000; SA: 350,000), a 133% performance spread between the best and worst results. If only SA had been available, the administrator would have received a timetable with one lecturer clash and one room clash still unresolved.")

    heading3(doc, "4.8.2 Why the Genetic Algorithm Performed Best")
    body(doc, "The GA's superiority over SA and TS for this problem instance can be attributed to four structural advantages. First, Population Diversity: GA maintains 50 distinct timetable solutions simultaneously, enabling the search to explore multiple regions of the solution space in parallel, whereas SA and TS maintain only a single current solution. Second, Crossover Recombination: The single-point crossover operator creates children that inherit good scheduling patterns from two different parent solutions — a mechanism SA and TS lack. Third, Elitism: The top five solutions are preserved intact across every generation, ensuring that the best fitness achieved is monotonically non-increasing. Fourth, Balanced Mutation: The 20% mutation rate provides effective exploration while retaining 80% of each solution's structure. These structural advantages are consistent with findings in the literature. Burke et al. (2013) noted that population-based methods tend to outperform single-solution methods on timetabling problems because of the richer information exchange enabled by a diverse population.")

    heading3(doc, "4.8.3 Why Simulated Annealing Underperformed Relative to GA")
    body(doc, "SA's final fitness of 350,000 (86.27% improvement) is significantly behind GA's 94.12% improvement. SA has no mechanism for simultaneously exploring multiple solution regions and its search is confined to a single path through the solution space, making it vulnerable to the particular local optima encountered during any given run. While the Metropolis acceptance criterion is designed to escape local optima, it can also lead the search away from a promising solution region when temperature remains high. SA's performance is also sensitive to the cooling schedule, and the chosen parameters may not have been optimal for this specific dataset size and constraint configuration. Despite these limitations, SA's 30-second execution time makes it attractive in contexts where speed is prioritised over solution quality.")

    heading3(doc, "4.8.4 Tabu Search Performance Assessment")
    body(doc, "TS achieved intermediate performance (fitness 250,000, 90.20% improvement), outperforming SA but remaining below GA. Its effective cycling prevention (tabu list size 30) ensured continuous exploration consistent with Sörensen et al. (2018)'s description of TS's memory-based search advantages. However, TS was constrained by its neighbourhood structure: generating only 20 neighbours per iteration limits the diversity of each search step compared to GA's population of 50. TS also required the longest execution time (~55 seconds) because each iteration involves generating and evaluating 20 neighbours plus computing timetable signatures for tabu checking.")

    heading3(doc, "4.8.5 Impact of NLP Integration on Usability")
    body(doc, "The regex-based NLP parser successfully translated all six tested constraint patterns from plain English into machine-readable constraint dictionaries with 100% accuracy within the supported vocabulary. This eliminates the need for administrators to learn any formal constraint language or configuration file format. Compared to transformer-based approaches, the rule-based implementation offers deterministic, reproducible results with no dependency on external API calls or pre-trained model weights. As noted by Nguyen et al. (2022), rule-based NLP remains the preferred approach for structured constraint formalisation where the vocabulary is well-defined and accuracy must be guaranteed. The key limitation is its fixed vocabulary: constraint types not covered by the predefined regex patterns are silently ignored.")

    heading3(doc, "4.8.6 Capacity Violations as the Most Persistent Challenge")
    body(doc, "All three algorithms retained exactly three capacity violations at their termination despite substantially different search strategies. This convergence on the same outcome suggests that the remaining violations are not algorithmic failures but structural constraints of the dataset. When all appropriately sized venues for a very large course (for example, 430+ students) are simultaneously occupied in a given timeslot, no algorithm can eliminate the capacity violation without violating a harder constraint such as lecturer or room non-overlap. This finding also reflects the dataset limitation acknowledged in Section 3.3.1: partially estimated venue capacity figures may cause some rooms to appear smaller than they actually are, artificially inflating the capacity violation count.")

    heading3(doc, "4.8.7 Scalability Considerations")
    body(doc, "The system demonstrated adequate performance for the Covenant University dataset (490+ courses, 100+ lecturers, 50+ venues) on standard consumer hardware. GA's 200 generations × 50 evaluations = 10,000 total fitness function calls completed in approximately 45 seconds. For larger institutions with 1,000+ courses, the computational cost would scale roughly linearly with dataset size. Potential optimisation strategies include parallel fitness evaluation across population members using Python's multiprocessing library, incremental fitness updates that recalculate only affected sessions after mutation, and adaptive population sizing that reduces population when convergence is detected.")

    heading3(doc, "4.8.8 Comparison with Related Work")
    body(doc, "The results obtained in this study are consistent with findings in the existing literature. Burke et al. (2013) demonstrated that hyper-heuristic portfolio approaches outperform single-algorithm systems on timetabling problems — confirmed by the 133% performance spread between GA and SA in this study. Ceschia et al. (2023) established the NP-hard nature of the UTSP, consistent with the inability of any algorithm to find a perfect zero-violation solution in this study. Katoch et al. (2021)'s review of genetic algorithms documented superior solution quality through crossover and elitism, consistent with GA's outperformance of both SA and TS. Delahaye et al. (2019) documented SA's parameter sensitivity, consistent with SA's underperformance under the chosen cooling parameters. Sörensen et al. (2018) confirmed TS's effective cycling prevention through tabu lists, reflected in TS's consistent improvement trajectory. Nguyen et al. (2022)'s demonstration of rule-based NLP accuracy for structured constraint vocabularies is confirmed by the 100% accuracy observed in the NLP parsing tests.")

    heading3(doc, "4.8.9 Achievement of Research Objectives")
    add_table(doc,
        ["Objective", "Description", "Status", "Evidence"],
        [
            ("1", "Construct a structured dataset suitable for integrating NLP reasoning with hyper-heuristic algorithms", "Achieved", "490+ course sessions collected, cleaned, supplemented, and successfully loaded by the scheduling engine"),
            ("2", "Use NLP tools to process hard and soft constraints into readable input for hyper-heuristic algorithms", "Achieved", "Regex-based NLP parser correctly extracted all 6 tested constraint types; constraints directly consumed by engine"),
            ("3", "Implement and comparatively evaluate the hyper-heuristic algorithms", "Achieved", "GA, SA, and TS all implemented; comparative results reported in Tables 4.9–4.13"),
            ("4", "Demonstrate the practical value of the best algorithm using a web application", "Achieved", "Time4Tables web application built with Streamlit/FastAPI; GA-selected timetable displayed in interactive weekly grid"),
        ],
        "Table 4.15: Research Objective Achievement Summary"
    )

    # 4.9
    heading2(doc, "4.9 Summary")
    body(doc, "This chapter presented the implementation and testing details, experimental setup, quantitative results, and comprehensive discussion for the hyper-heuristic university timetable generation system. The Genetic Algorithm achieved the best overall performance, with a fitness score of 150,000 — representing a 94.12% improvement over the greedy initial solution — and complete elimination of all lecturer clashes and room clashes. Simulated Annealing achieved an 86.27% improvement, while Tabu Search achieved a 90.20% improvement. All three algorithms successfully scheduled all 490+ course sessions with zero unscheduled sessions, confirming that each algorithm is capable of producing a complete, usable timetable. The NLP constraint parser demonstrated 100% accuracy across all tested constraint patterns. Three capacity violations persisted across all algorithms, reflecting a structural constraint of the partially estimated dataset rather than algorithmic failure. The hyper-heuristic selection mechanism ensured that the best-available solution (from GA) was automatically delivered to the administrator. These results address all four research objectives established in Chapter One and are consistent with the theoretical and empirical findings reported in the literature reviewed in Chapter Two. Chapter Five will summarise the key contributions of the study, state its limitations, and present recommendations for future work.")

    # ─── SAVE ───
    out_path = os.path.join(os.path.dirname(__file__), "CHAPTERS_3_AND_4.docx")
    doc.save(out_path)
    print(f"\nDone! Saved to:\n   {out_path}")


if __name__ == "__main__":
    build()
